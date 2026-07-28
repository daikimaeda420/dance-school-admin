"""Create the customer-facing RIZBO initial setup guide (DOCX)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "guides" / "rizbo-school-setup-guide.docx"

ORANGE = RGBColor(254, 97, 71)
BLUE = RGBColor(37, 99, 235)
INK = RGBColor(23, 32, 51)
MUTED = RGBColor(71, 85, 105)
LIGHT_ORANGE = "FFF3EE"
LIGHT_BLUE = "EFF6FF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_font(run, size=10.5, color=INK, bold=False) -> None:
    run.font.name = "YuGothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "YuGothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "YuGothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "YuGothic")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_paragraph(doc, text="", *, size=10.5, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.28
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 12.5, color=BLUE if level == 1 else INK, bold=True)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    label_run = p.add_run(f"{label}  ")
    set_font(label_run, size=10, color=BLUE if fill == LIGHT_BLUE else ORANGE, bold=True)
    body_run = p.add_run(text)
    set_font(body_run, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_step(doc, number: int, title: str, purpose: str, actions: list[str], check: str):
    add_heading(doc, f"ステップ {number}　{title}", level=1)
    add_paragraph(doc, purpose, color=MUTED, after=7)
    for action in actions:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        set_font(p.add_run(action), size=10.5, color=INK)
    add_callout(doc, "確認", check, LIGHT_ORANGE)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("RIZBO 初期設定ガイド  |  ")
    set_font(run, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "YuGothic"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "YuGothic")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "YuGothic")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "YuGothic")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("RIZBO 初期設定ガイド")
    set_font(run, size=25, color=INK, bold=True)

    add_paragraph(
        doc,
        "Q&A・相性診断・予約フォームを整え、スクールサイトへ公開するまでの手順です。",
        size=12,
        color=MUTED,
        after=16,
    )

    info = doc.add_table(rows=1, cols=2)
    set_table_geometry(info, [2100, 7260])
    for i, (label, value) in enumerate((("対象", "RIZBOを導入するスクールご担当者"), ("所要時間の目安", "初回設定：20〜40分"))):
        cell = info.cell(0, i)
        set_cell_shading(cell, "F8FAFC")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        label_run = p.add_run(f"{label}\n")
        set_font(label_run, size=8.5, color=MUTED, bold=True)
        value_run = p.add_run(value)
        set_font(value_run, size=10.5, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    add_heading(doc, "はじめる前に", level=1)
    add_paragraph(doc, "設定を始める前に、次の2点をご用意ください。")
    for text in (
        "スクールサイトを編集できる権限（自社・制作会社のどちらが編集するかも確認）",
        "体験予約を受け取るメールアドレス",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(text), size=10.5, color=INK)

    add_step(
        doc,
        1,
        "Q&Aを整える",
        "よくある質問と回答を登録し、サイト訪問者が迷わず次の行動へ進めるようにします。",
        [
            "管理画面の「コンテンツ設定」→「Q&A編集」を開きます。",
            "料金・体験・初心者・アクセス・予約に関する質問から追加します。",
            "各質問に、短く具体的な回答と必要に応じたリンクを設定して保存します。",
        ],
        "質問と回答が空欄になっていないか、料金や予約に関する案内が入っているかを確認します。",
    )
    add_step(
        doc,
        2,
        "Q&Aの完成度を確認する",
        "公開前の設定漏れや申込導線の不足を、チェック形式で確認できます。",
        [
            "「Q&A完成度」を開きます。",
            "「未設定」「要確認」と表示される項目を、案内に沿って修正します。",
            "すべての項目を確認したら、次の診断設定へ進みます。",
        ],
        "完成度だけでなく、各項目の説明を読み、実際のスクール情報と合っているかも確認します。",
    )
    add_step(
        doc,
        3,
        "診断コンテンツを設定する",
        "相性診断で提案するコース・講師・スケジュールを登録します。",
        [
            "「診断編集」を開き、まず校舎とコースを登録します。",
            "ジャンル・年代／ライフスタイル・講師・スケジュールを順に設定します。",
            "コースと講師の説明は、初めての方にも伝わる内容にします。",
        ],
        "診断の回答に対して、紹介したいコースがきちんと候補に出る状態を目指します。",
    )
    add_step(
        doc,
        4,
        "予約フォームを設定する",
        "診断後の体験予約を受け取れるよう、フォームと通知先を確認します。",
        [
            "「診断編集」→「フォーム」を開きます。",
            "氏名・メールアドレスなど、予約に必要な項目だけを必須にします。",
            "送信先メールアドレスを確認し、テスト送信で受信できるか確認します。",
        ],
        "入力項目が多すぎると離脱につながるため、最初は必要最小限の項目がおすすめです。",
    )
    add_step(
        doc,
        5,
        "プレビューで体験する",
        "サイト訪問者と同じ流れを実際に操作して、公開前に見え方を確認します。",
        [
            "左メニュー下部の「チャットプレビュー」でQ&Aを確認します。",
            "「診断プレビュー」で、診断開始からフォーム送信まで操作します。",
            "スマートフォンでも表示と入力のしやすさを確認します。",
        ],
        "質問の順番、選択肢の分かりやすさ、フォームの入力負担を重点的に見てください。",
    )
    add_step(
        doc,
        6,
        "サイトに設置して公開する",
        "RIZBOの設置コードをスクールサイトへ追加し、公開後に動作を確認します。",
        [
            "管理画面の「ホーム」を開き、「設置コード」をコピーします。",
            "スクールサイトのHTMLにコードを追加します。制作会社が管理している場合は、コピーしたコードを共有します。",
            "公開後、実際のサイトでチャットと診断が表示されることを確認します。",
        ],
        "設置後は「ホーム」で設置サイトUU、「AIコンサル・分析」で診断完了率・予約率を確認できます。",
    )

    add_heading(doc, "公開後の運用", level=1)
    add_paragraph(doc, "公開して終わりではなく、月に一度は成果を見ながら改善を続けることが大切です。")
    for text in (
        "ホーム：設置サイトUU、Q&A・診断の利用状況を確認",
        "AIコンサル・分析：離脱ポイント、需要トレンド、改善提案を確認",
        "月次KPI・売上シミュレーション：予約数と推定売上を振り返り",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(text), size=10.5, color=INK)

    add_callout(doc, "困ったときは", "管理画面の「設定ガイド」に戻ると、各設定画面へのリンクと確認の順番をいつでも確認できます。")
    add_paragraph(doc, "設定ガイドURL：https://rizbo.dansul.jp/admin/setup-guide", size=9.5, color=MUTED, after=0)

    add_page_number(section)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
